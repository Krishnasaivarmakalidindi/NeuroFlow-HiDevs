import docx
from typing import List

from monitoring.tracing import trace_sync

@trace_sync("ingestion.extract.docx")
def extract_docx(file_path: str) -> List[ExtractedPage]:
    doc = docx.Document(file_path)
    pages = []
    
    current_heading = None
    content_blocks = []
    
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            # It's a header
            level = f"h{para.style.name.split(' ')[-1]}"
            current_heading = {"level": level, "section": para.text}
            content_blocks.append(para.text)
        elif para.text.strip():
            # Paragraph
            content_blocks.append(para.text)
            
    # For simplification, DOCX doesn't have strict pages. We group everything as page 1.
    if content_blocks:
        pages.append(ExtractedPage(
            page_number=1,
            content="\n".join(content_blocks),
            content_type="text",
            metadata=current_heading if current_heading else {}
        ))
        
    # Extract tables
    for i, table in enumerate(doc.tables):
        md_table = []
        for row in table.rows:
            cleaned_row = [cell.text.replace('\n', ' ') if cell.text else "" for cell in row.cells]
            md_table.append("| " + " | ".join(cleaned_row) + " |")
        
        if md_table:
            headers = md_table[0]
            separator = "|---" * headers.count("|") + "|"
            md_table.insert(1, separator)
            
            pages.append(ExtractedPage(
                page_number=1,
                content="\n".join(md_table),
                content_type="table",
                metadata={"table_index": i}
            ))

    return pages
